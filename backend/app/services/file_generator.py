# backend/app/services/file_generator.py

import os
from docx import Document

def replace_text(paragraph, key, value):
    if key in paragraph.text:
        for run in paragraph.runs:
            if key in run.text:
                run.text = run.text.replace(key, str(value).upper())

def generate_documents(files, output_dir, data):

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for template in files:
        doc = Document(template)

        for key, value in data.items():
            for paragraph in doc.paragraphs:
                replace_text(paragraph, key, value)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text(paragraph, key, value)

        output_path = os.path.join(output_dir, os.path.basename(template))
        doc.save(output_path)

    return output_dir